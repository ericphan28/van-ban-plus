"""
Tạo file Word: Yêu cầu tính năng AI — Góc nhìn Cán bộ
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ═══ Page Setup ═══
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)

# ═══ Styles ═══
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.3

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_para(text, bold=False, italic=False, size=13, color=None, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    else:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '003366',
            qn('w:val'): 'clear'
        })
        shading.append(shading_elm)
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        # Alternate row shading
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                shading = table.rows[r_idx + 1].cells[c_idx]._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F0F4F8',
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elm)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_box(title, content_lines, color_hex='E8F5E9', border_color='4CAF50'):
    """Add a colored box with title and bullet content"""
    p = doc.add_paragraph()
    run = p.add_run(f"  {title}")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 100, 0)
    for line in content_lines:
        p = doc.add_paragraph()
        run = p.add_run(f"    {line}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(2)

def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("─" * 70)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.size = Pt(8)

# ═══════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

add_para("CÔNG TY TNHH GIA KIỆM SỐ", bold=True, size=14, 
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0, 51, 102))
add_para("giakiemso.com", italic=True, size=12, 
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100))

doc.add_paragraph()

add_para("YÊU CẦU TÍNH NĂNG AI", bold=True, size=22, 
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0, 51, 102))
add_para("GÓC NHÌN CÁN BỘ", bold=True, size=16, 
         align=WD_ALIGN_PARAGRAPH.CENTER, color=(0, 100, 150))

doc.add_paragraph()

add_para("Ứng dụng: VanBanPlus — Quản lý văn bản hành chính thông minh", 
         size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Phiên bản: 1.0  •  Ngày: 13/02/2026", 
         size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100))

doc.add_paragraph()
doc.add_paragraph()

add_table(
    ["Thông tin", "Chi tiết"],
    [
        ["Người yêu cầu", "Cán bộ Văn phòng — Thống kê UBND xã"],
        ["Đối tượng sử dụng", "Cán bộ, công chức cấp xã/huyện/tỉnh"],
        ["Mục tiêu", "Giảm thời gian xử lý VB từ 4-6 giờ/ngày xuống 1-2 giờ/ngày"],
        ["AI Engine", "Google Gemini 2.5 Flash"],
        ["Nền tảng", "Windows Desktop (WPF, .NET 9)"],
    ],
    col_widths=[5, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════
# MỤC LỤC
# ═══════════════════════════════════════════════
add_heading_styled("MỤC LỤC", level=1)
toc_items = [
    "1. AI Tạo Văn Bản — Soạn thảo tự động từ mẫu",
    "2. AI Scan OCR — Đọc ảnh/PDF thành dữ liệu",
    "3. AI Kiểm Tra — Soát lỗi chính tả, thể thức, văn phong",
    "4. AI Tham Mưu — Đề xuất xử lý văn bản đến",
    "5. AI Tóm Tắt — Tóm tắt văn bản dài thành 10 mục",
    "6. AI Báo Cáo — Viết báo cáo định kỳ từ số liệu",
    "Tổng hợp hiệu quả"
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

doc.add_page_break()

# ═══════════════════════════════════════════════
# 1. AI TẠO VĂN BẢN
# ═══════════════════════════════════════════════
add_heading_styled("1. AI TẠO VĂN BẢN", level=1)
add_para("Soạn thảo văn bản hành chính tự động từ mẫu có sẵn bằng trí tuệ nhân tạo.", italic=True, color=(80, 80, 80))

add_heading_styled("Nỗi đau hiện tại", level=2)
add_para("Mỗi ngày phải soạn 3-5 văn bản. Mỗi văn bản mất 45-90 phút vì phải:")
add_bullet("Mở file Word cũ → copy → sửa → quên đổi ngày/tên → bị lãnh đạo trả lại")
add_bullet("Không nhớ thể thức đúng theo Nghị định 30/2020/NĐ-CP")
add_bullet("Viết đi viết lại phần mở đầu, căn cứ pháp lý")

add_heading_styled("Tính năng cần", level=2)
add_para("Chọn loại văn bản → Nhập thông tin cốt lõi → AI tạo bản nháp hoàn chỉnh.", bold=True)

add_heading_styled("Ví dụ minh họa", level=2)
add_para("Tình huống: Chủ tịch xã giao soạn Công văn mời họp Ban chỉ đạo phòng chống bão lụt.", bold=True, color=(0, 80, 0))

add_para("TRƯỚC KHI CÓ AI (45 phút):", bold=True, color=(180, 0, 0))
add_bullet("Tìm file CV mời họp cũ trong máy → 10 phút")
add_bullet("Copy sang file mới, sửa nội dung → 20 phút")
add_bullet("Sửa lại thể thức (quên đổi số, ngày, nơi nhận) → 10 phút")
add_bullet("In ra, lãnh đạo phát hiện sai căn cứ → sửa thêm 5 phút")

add_para("SAU KHI CÓ AI (5 phút):", bold=True, color=(0, 120, 0))
add_bullet('Chọn mẫu "Công văn mời họp"')
add_bullet('Nhập: Nội dung = "Triển khai phòng chống bão số 3", Thời gian = "14h ngày 15/02/2026"')
add_bullet("AI tạo CV hoàn chỉnh: đúng thể thức, đúng căn cứ, đúng format")
add_bullet("Xem lại → Lưu → Xuất Word → In")

add_para("▶ Tiết kiệm: ~40 phút/văn bản × 4 văn bản/ngày = 160 phút/ngày", bold=True, color=(0, 100, 0))

add_separator()

# ═══════════════════════════════════════════════
# 2. AI SCAN OCR
# ═══════════════════════════════════════════════
add_heading_styled("2. AI SCAN OCR", level=1)
add_para("Đọc ảnh chụp / file PDF scan → trích xuất tự động 14 trường dữ liệu.", italic=True, color=(80, 80, 80))

add_heading_styled("Nỗi đau hiện tại", level=2)
add_para("Mỗi tuần nhận 20-30 văn bản giấy từ huyện, tỉnh. Phải:")
add_bullet("Ngồi đọc từng tờ → gõ lại số VB, ngày, trích yếu → 10-15 phút/văn bản")
add_bullet("Gõ sai số, sai ngày → tra cứu sau không tìm thấy")
add_bullet("Văn bản chất đống, không kịp nhập → bị nhắc nhở")

add_heading_styled("Tính năng cần", level=2)
add_para("Chụp ảnh/scan văn bản → AI tự đọc → trích xuất đầy đủ thông tin → lưu vào hệ thống.", bold=True)

add_heading_styled("14 trường AI tự động trích xuất", level=2)
add_table(
    ["#", "Trường", "Ví dụ"],
    [
        ["1", "Số văn bản", "456/QĐ-UBND"],
        ["2", "Trích yếu", "V/v phân bổ kinh phí xây dựng NTM"],
        ["3", "Loại văn bản", "Quyết định"],
        ["4", "Ngày ban hành", "10/02/2026"],
        ["5", "Cơ quan ban hành", "UBND huyện XYZ"],
        ["6", "Người ký", "Nguyễn Văn A"],
        ["7", "Nội dung", "Toàn văn nội dung văn bản"],
        ["8", "Nơi nhận", "Sở Tài chính, UBND các xã..."],
        ["9", "Căn cứ pháp lý", "Luật Ngân sách nhà nước 2015..."],
        ["10", "Hướng văn bản", "Đến"],
        ["11", "Lĩnh vực", "Kinh tế"],
        ["12", "Địa danh", "Biên Hòa"],
        ["13", "Chức danh ký", "CHỦ TỊCH"],
        ["14", "Thẩm quyền ký", "TM. UBND"],
    ],
    col_widths=[1, 4, 12]
)

add_heading_styled("Ví dụ minh họa", level=2)
add_para("Tình huống: Nhận QĐ số 456/QĐ-UBND ngày 10/02/2026 của UBND huyện về phân bổ kinh phí.", bold=True, color=(0, 80, 0))

add_para("TRƯỚC KHI CÓ AI (15 phút):", bold=True, color=(180, 0, 0))
add_bullet("Đọc QĐ giấy → ghi ra giấy nháp → nhập thủ công 10+ trường")
add_bullet('Gõ nhầm "456" thành "465" → sau này tìm không ra')
add_bullet("Quên nhập căn cứ pháp lý → thiếu thông tin khi cần tra cứu")

add_para("SAU KHI CÓ AI (2 phút):", bold=True, color=(0, 120, 0))
add_bullet("Chụp ảnh QĐ bằng điện thoại → gửi về máy tính")
add_bullet('Nhấn "AI Scan OCR" → chọn ảnh → AI trích xuất tất cả 14 trường')
add_bullet("Kiểm tra nhanh → Lưu — không sai sót")

add_para("▶ Tiết kiệm: ~13 phút/VB × 25 VB/tuần = 325 phút/tuần (~5.4 giờ)", bold=True, color=(0, 100, 0))

add_separator()

# ═══════════════════════════════════════════════
# 3. AI KIỂM TRA VĂN BẢN
# ═══════════════════════════════════════════════
add_heading_styled("3. AI KIỂM TRA VĂN BẢN", level=1)
add_para("Soát lỗi chính tả, văn phong, thể thức theo NĐ 30/2020 trước khi trình ký.", italic=True, color=(80, 80, 80))

add_heading_styled("Nỗi đau hiện tại", level=2)
add_para("Soạn xong văn bản, in ra trình ký → lãnh đạo phát hiện:")
add_bullet('Sai chính tả ("khẩn trương" → "khẩn chương")')
add_bullet("Căn cứ pháp lý đã hết hiệu lực")
add_bullet('Thiếu nơi nhận "Lưu VT"')
add_bullet("Không đúng thể thức (quên Quốc hiệu, sai format số)")
add_para("→ Trả lại sửa 2-3 lần, mất uy tín + mất thời gian cả cán bộ lẫn lãnh đạo.", bold=True, color=(180, 0, 0))

add_heading_styled("Tính năng cần", level=2)
add_para("Trước khi trình ký → AI kiểm tra toàn bộ 8 khía cạnh → liệt kê lỗi + gợi ý sửa.", bold=True)

add_heading_styled("8 khía cạnh AI kiểm tra", level=2)
add_table(
    ["#", "Khía cạnh", "Kiểm tra gì"],
    [
        ["1", "Chính tả", "Lỗi typo, viết hoa tiếng Việt"],
        ["2", "Văn phong", "Đúng ngôn ngữ hành chính không"],
        ["3", "Xung đột nội dung", "Các đoạn mâu thuẫn nhau"],
        ["4", "Logic & cấu trúc", "Đánh số liên tục, tham chiếu hợp lệ"],
        ["5", "Thiếu thành phần", "Thiếu căn cứ, nơi nhận theo loại VB"],
        ["6", "Nội dung mơ hồ", "Chủ thể, deadline, số liệu không rõ"],
        ["7", "Đề xuất cải thiện", "Gợi ý viết tốt hơn"],
        ["8", "Thể thức NĐ 30/2020", "Quốc hiệu, tiêu ngữ, số/ký hiệu, chữ ký"],
    ],
    col_widths=[1, 4.5, 11.5]
)

add_heading_styled("Ví dụ minh họa", level=2)
add_para("Tình huống: Soạn Tờ trình đề nghị UBND huyện hỗ trợ kinh phí sửa chữa trường học.", bold=True, color=(0, 80, 0))

add_para("TRƯỚC KHI CÓ AI (bị trả lại 3 lần = 90 phút):", bold=True, color=(180, 0, 0))
add_bullet('Lần 1: Lãnh đạo phát hiện "Thiếu căn cứ Luật Ngân sách nhà nước" → trả lại')
add_bullet('Lần 2: "Nơi nhận thiếu Phòng TC-KH huyện" → trả lại')
add_bullet('Lần 3: "Viết sai UBND thành UNBD" → trả lại')

add_para("SAU KHI CÓ AI (5 phút — duyệt ngay lần đầu):", bold=True, color=(0, 120, 0))
add_bullet('🔴 Lỗi nghiêm trọng: Thiếu căn cứ "Luật Ngân sách nhà nước 2015"', bold_prefix="")
add_bullet('🔴 Lỗi nghiêm trọng: Nơi nhận thiếu "Phòng TC-KH huyện"', bold_prefix="")
add_bullet('🟡 Cảnh báo: Lỗi chính tả "UNBD" → "UBND" ở đoạn 3', bold_prefix="")
add_bullet('🟢 Gợi ý: Thêm số liệu cụ thể về mức kinh phí đề nghị', bold_prefix="")
add_para("→ Sửa tất cả → trình ký → duyệt ngay lần đầu!", bold=True, color=(0, 120, 0))

add_para("▶ Tiết kiệm: ~60 phút mỗi VB bị trả lại. Giảm 90% tỷ lệ văn bản bị trả.", bold=True, color=(0, 100, 0))

add_separator()

# ═══════════════════════════════════════════════
# 4. AI THAM MƯU XỬ LÝ
# ═══════════════════════════════════════════════
add_heading_styled("4. AI THAM MƯU XỬ LÝ", level=1)
add_para("Phân tích văn bản đến → đề xuất: ai xử lý, deadline, cần trả lời gì, rủi ro gì.", italic=True, color=(80, 80, 80))

add_heading_styled("Nỗi đau hiện tại", level=2)
add_para("Nhận văn bản từ cấp trên, cán bộ không biết:")
add_bullet("Ai xử lý? Chủ tịch hay Phó CT?")
add_bullet("Deadline bao lâu? 5 ngày hay 10 ngày?")
add_bullet("Cần trả lời bằng loại văn bản nào?")
add_bullet("Có liên quan đến văn bản nào trước đó?")
add_para("→ Hỏi đồng nghiệp, hỏi lãnh đạo → mất 30-60 phút mỗi VB phức tạp. Hoặc xử lý sai → bị nhắc nhở, trễ hạn.", bold=True, color=(180, 0, 0))

add_heading_styled("Tính năng cần", level=2)
add_para("Nhận VB đến → AI đọc hiểu → đề xuất xử lý theo 15 chiều phân tích.", bold=True)

add_heading_styled("Ví dụ minh họa", level=2)
add_para('Tình huống: Nhận CV số 789/UBND-NV ngày 12/02/2026 của UBND huyện yêu cầu "Báo cáo CCHC năm 2025 trước ngày 20/02/2026".', bold=True, color=(0, 80, 0))

add_para("TRƯỚC KHI CÓ AI (45 phút):", bold=True, color=(180, 0, 0))
add_bullet("Đọc CV → không chắc thuộc lĩnh vực ai phụ trách → hỏi VP → 15 phút")
add_bullet("Không biết cần trả lời bằng Báo cáo hay Công văn → hỏi đồng nghiệp → 10 phút")
add_bullet("Không nhớ năm trước làm thế nào → tìm file cũ → 15 phút")
add_bullet("Suýt quên deadline 20/02")

add_para("SAU KHI CÓ AI (3 phút) — Kết quả phân tích:", bold=True, color=(0, 120, 0))
add_table(
    ["Mục phân tích", "Kết quả AI"],
    [
        ["Tóm tắt", "Huyện yêu cầu báo cáo CCHC năm 2025"],
        ["Mức khẩn", "🟡 Khẩn (còn 8 ngày)"],
        ["Deadline", "20/02/2026 (trích từ CV)"],
        ["Người xử lý", "Phó CT phụ trách Văn xã, phối hợp VP-TK"],
        ["Thẩm quyền ký", "Chủ tịch UBND xã"],
        ["Cần trả lời", "Có — bằng Báo cáo"],
        ["Dự thảo phản hồi", "I. Kết quả CCHC: (1) Thủ tục HC, (2) Tổ chức bộ máy..."],
        ["Căn cứ pháp lý", "NQ 76/NQ-CP, QĐ 468/QĐ-TTg về CCHC"],
        ["Cảnh báo rủi ro", "⚠ Trễ hạn sẽ bị trừ điểm thi đua đơn vị"],
    ],
    col_widths=[4.5, 12.5]
)

add_para("▶ Tiết kiệm: ~40 phút/VB phức tạp. Không bao giờ trễ hạn vì quên.", bold=True, color=(0, 100, 0))

add_separator()

# ═══════════════════════════════════════════════
# 5. AI TÓM TẮT VĂN BẢN
# ═══════════════════════════════════════════════
add_heading_styled("5. AI TÓM TẮT VĂN BẢN", level=1)
add_para("Tóm tắt văn bản dài (Nghị định, Thông tư) thành 10 mục có cấu trúc.", italic=True, color=(80, 80, 80))

add_heading_styled("Nỗi đau hiện tại", level=2)
add_para("Nhận Nghị định 50 trang, Thông tư 30 trang → phải đọc hết để:")
add_bullet("Nắm nội dung chính để báo cáo lãnh đạo")
add_bullet("Tìm điều khoản liên quan đến xã")
add_bullet("Trích dẫn cho văn bản đang soạn")
add_para("→ Đọc 1 Nghị định mất 2-3 giờ, mà mỗi tuần nhận 5-10 VB dài.", bold=True, color=(180, 0, 0))

add_heading_styled("Tính năng cần", level=2)
add_para("AI đọc toàn bộ → tóm tắt 10 mục: nội dung chính, đối tượng, thời hạn, số liệu, tác động.", bold=True)

add_heading_styled("Ví dụ minh họa", level=2)
add_para("Tình huống: Nhận Nghị định 35 trang về quản lý đất đai, cần báo cáo CT xã trong buổi giao ban sáng mai.", bold=True, color=(0, 80, 0))

add_para("TRƯỚC KHI CÓ AI (3 giờ):", bold=True, color=(180, 0, 0))
add_bullet("Đọc 35 trang → gạch chân phần quan trọng → 2 giờ")
add_bullet("Tóm tắt ra giấy → 30 phút")
add_bullet('Vẫn bỏ sót 2 điều khoản quan trọng. Lãnh đạo hỏi "Điều 15 nói gì?" → không nhớ')

add_para("SAU KHI CÓ AI (5 phút) — Kết quả tóm tắt:", bold=True, color=(0, 120, 0))
add_table(
    ["Mục", "Nội dung AI tóm tắt"],
    [
        ["Tóm tắt", "NĐ quy định về quyền sử dụng đất, chuyển mục đích, cấp GCN..."],
        ["Đối tượng", "UBND cấp xã, huyện, tỉnh; Hộ gia đình, tổ chức"],
        ["Nội dung chính", "① Điều 5-8: Thu hồi đất  ② Điều 12: Cấp GCN  ③ Điều 15: Chuyển mục đích  ④ Điều 20-22: Bồi thường"],
        ["Thời hạn", "Có hiệu lực từ 01/07/2026"],
        ["Số liệu", "Mức bồi thường tối thiểu: 1.2 lần giá đất"],
        ["Tác động", "Xã cần: cập nhật quy trình, tập huấn cán bộ địa chính"],
    ],
    col_widths=[4, 13]
)

add_para("▶ Tiết kiệm: ~2.5 giờ/VB dài. Không bỏ sót nội dung quan trọng.", bold=True, color=(0, 100, 0))

add_separator()

# ═══════════════════════════════════════════════
# 6. AI BÁO CÁO ĐỊNH KỲ
# ═══════════════════════════════════════════════
add_heading_styled("6. AI BÁO CÁO ĐỊNH KỲ", level=1)
add_para("Viết báo cáo định kỳ từ số liệu thô, tự tính % tăng/giảm so với kỳ trước.", italic=True, color=(80, 80, 80))

add_heading_styled("Nỗi đau hiện tại", level=2)
add_para("Mỗi tháng phải làm 4-6 báo cáo (KT-XH, CCHC, Nội vụ, ANTT...). Mỗi báo cáo:")
add_bullet("Thu thập số liệu từ các bộ phận → 1 giờ")
add_bullet("Viết phần nhận xét, đánh giá, so sánh kỳ trước → 2-3 giờ")
add_bullet("Tính % tăng/giảm → hay sai số")
add_bullet("Sếp yêu cầu sửa văn phong → thêm 1 giờ")
add_para("→ Riêng viết báo cáo chiếm 2-3 ngày/tháng.", bold=True, color=(180, 0, 0))

add_heading_styled("Tính năng cần", level=2)
add_para("Nhập số liệu thô + chọn kỳ/lĩnh vực → AI viết báo cáo hoàn chỉnh 3 phần.", bold=True)

add_heading_styled("Ví dụ minh họa", level=2)
add_para("Tình huống: Làm Báo cáo KT-XH tháng 01/2026 cho UBND xã.", bold=True, color=(0, 80, 0))

add_para("TRƯỚC KHI CÓ AI (4 giờ):", bold=True, color=(180, 0, 0))
add_bullet("Thu thập số liệu → 1 giờ")
add_bullet('Mở BC tháng trước → copy → sửa → hay quên đổi "tháng 12" thành "tháng 01"')
add_bullet("Tính tay: 2.5 tỷ / 2.1 tỷ = tăng 19% → 30 phút (hay sai)")
add_bullet('Viết nhận xét + phương hướng → 2.5 giờ')

add_para("SAU KHI CÓ AI (15 phút):", bold=True, color=(0, 120, 0))
add_bullet('Chọn: Kỳ = "Tháng", Lĩnh vực = "Kinh tế - Xã hội"')
add_bullet("Nhập số liệu thô: Thu NS 2.5 tỷ, Hộ nghèo giảm 3, GPMB 85%...")
add_bullet("Dán BC tháng 12/2025 (để AI so sánh)")
add_para("AI tự động tạo:", bold=True)
add_para('   "I. KẾT QUẢ THỰC HIỆN', italic=True)
add_para('   1. Thu ngân sách tháng 01/2026 đạt 2,5 tỷ đồng, tăng 19,05% so với tháng 12/2025...', italic=True)
add_para('   II. ĐÁNH GIÁ CHUNG — Tình hình KT-XH tiếp tục ổn định và tích cực...', italic=True)
add_para('   III. PHƯƠNG HƯỚNG THÁNG 02/2026 — Đẩy nhanh tiến độ GPMB 15% còn lại..."', italic=True)

add_para("▶ Tiết kiệm: ~3.5 giờ/báo cáo × 5 BC/tháng = 17.5 giờ/tháng (~2 ngày làm việc)", bold=True, color=(0, 100, 0))

doc.add_page_break()

# ═══════════════════════════════════════════════
# TỔNG HỢP HIỆU QUẢ
# ═══════════════════════════════════════════════
add_heading_styled("TỔNG HỢP HIỆU QUẢ", level=1)

add_heading_styled("So sánh thời gian xử lý", level=2)
add_table(
    ["Tính năng", "Trước AI", "Sau AI", "Tiết kiệm"],
    [
        ["Soạn 1 văn bản", "45-90 phút", "5-10 phút", "~40-80 phút"],
        ["Nhập 1 VB giấy", "10-15 phút", "2 phút", "~10 phút"],
        ["Kiểm tra 1 VB", "30-90 phút (sửa 2-3 lần)", "5 phút (sửa 1 lần)", "~60 phút"],
        ["Tham mưu 1 VB đến", "30-60 phút", "3 phút", "~40 phút"],
        ["Tóm tắt 1 VB dài", "2-3 giờ", "5 phút", "~2.5 giờ"],
        ["Làm 1 BC định kỳ", "3-4 giờ", "15 phút", "~3.5 giờ"],
    ],
    col_widths=[4.5, 4.5, 4, 4]
)

doc.add_paragraph()
add_heading_styled("Ước tính hiệu quả 1 tháng cho 1 cán bộ VP-TK", level=2)
add_table(
    ["Công việc", "Số lượng/tháng", "Giờ tiết kiệm"],
    [
        ["Soạn văn bản", "~60 VB", "40 giờ"],
        ["Nhập VB giấy", "~80 VB", "13 giờ"],
        ["Kiểm tra VB", "~30 VB", "30 giờ"],
        ["Tham mưu VB đến", "~40 VB", "27 giờ"],
        ["Tóm tắt VB dài", "~10 VB", "25 giờ"],
        ["BC định kỳ", "~5 BC", "17 giờ"],
        ["TỔNG CỘNG", "", "~152 giờ/tháng (~19 ngày)"],
    ],
    col_widths=[6, 4.5, 6.5]
)

doc.add_paragraph()
add_para("KẾT LUẬN", bold=True, size=14, color=(0, 51, 102))
add_para(
    "AI không thay thế cán bộ mà giúp cán bộ hoàn thành công việc nhanh gấp 5-10 lần, "
    "giảm sai sót, không trễ hạn. Thời gian tiết kiệm được dùng cho công việc cần tư duy: "
    "tiếp dân, giải quyết hồ sơ, đi cơ sở.",
    size=13
)

doc.add_paragraph()
add_separator()
add_para("Công ty TNHH Gia Kiệm Số — giakiemso.com", 
         italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(100, 100, 100))

# ═══ Save ═══
output_path = r"D:\AIVanBanCaNhan\REQUIREMENT_AI_CANBO.docx"
doc.save(output_path)
print(f"✅ Đã tạo file Word: {output_path}")
