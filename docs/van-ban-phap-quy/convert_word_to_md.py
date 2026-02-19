"""
Script chuyển đổi file Word (.docx) sang Markdown (.md)
Dùng để chuyển văn bản pháp quy gốc sang format Copilot đọc được.

Cách dùng:
    python convert_word_to_md.py "file.docx"
    python convert_word_to_md.py "file.docx" -o "output.md"
    python convert_word_to_md.py "folder/"          # convert tất cả .docx trong folder

Yêu cầu:
    pip install python-docx
"""

import sys
import os
import re
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Cần cài python-docx:")
    print("   pip install python-docx")
    sys.exit(1)


def docx_to_markdown(docx_path: str, output_path: str = None) -> str:
    """Chuyển file .docx sang Markdown text."""
    doc = Document(docx_path)
    lines = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        
        style_name = (para.style.name or "").lower()
        
        # Detect headings
        if "heading 1" in style_name or "title" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "heading 4" in style_name:
            lines.append(f"#### {text}")
        # Detect "Điều X." pattern → heading 3
        elif re.match(r'^Điều \d+[\.\:]', text):
            lines.append(f"### {text}")
        # Detect "Chương X" pattern → heading 2
        elif re.match(r'^Chương [IVXLCDM]+', text) or re.match(r'^CHƯƠNG [IVXLCDM]+', text):
            lines.append(f"## {text}")
        # Detect "Mục X" pattern
        elif re.match(r'^Mục \d+', text):
            lines.append(f"#### {text}")
        # Detect all-caps lines (likely section headers)
        elif text.isupper() and len(text) < 100:
            lines.append(f"## {text}")
        # List items
        elif text.startswith(("- ", "• ", "– ")):
            lines.append(f"- {text[2:]}")
        elif re.match(r'^\d+[\.\)]\s', text):
            lines.append(text)  # numbered list
        elif re.match(r'^[a-zđ][\.\)]\s', text):
            lines.append(f"- {text}")  # lettered list → bullet
        else:
            # Check for bold runs
            has_bold = any(run.bold for run in para.runs if run.text.strip())
            all_bold = all(run.bold for run in para.runs if run.text.strip()) if para.runs else False
            
            if all_bold and len(text) < 80:
                lines.append(f"**{text}**")
            else:
                lines.append(text)
    
    # Process tables
    for table in doc.tables:
        lines.append("")
        # Header row
        header = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        # Data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")
    
    markdown = "\n".join(lines)
    
    # Clean up multiple blank lines
    markdown = re.sub(r'\n{3,}', '\n\n', markdown)
    
    # Save if output path provided
    if output_path:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        print(f"✅ Đã chuyển: {docx_path} → {output_path}")
    
    return markdown


def process_folder(folder_path: str):
    """Chuyển tất cả .docx trong folder."""
    folder = Path(folder_path)
    docx_files = list(folder.glob("*.docx"))
    
    if not docx_files:
        print(f"⚠️ Không tìm thấy file .docx nào trong {folder_path}")
        return
    
    print(f"📂 Tìm thấy {len(docx_files)} file .docx")
    
    for docx_file in docx_files:
        output = docx_file.with_suffix('.md')
        docx_to_markdown(str(docx_file), str(output))
    
    print(f"\n🎉 Hoàn thành! Đã chuyển {len(docx_files)} file.")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    
    input_path = sys.argv[1]
    
    if os.path.isdir(input_path):
        process_folder(input_path)
    elif os.path.isfile(input_path):
        output = sys.argv[3] if len(sys.argv) > 3 and sys.argv[2] == "-o" else input_path.replace('.docx', '.md')
        docx_to_markdown(input_path, output)
    else:
        print(f"❌ Không tìm thấy: {input_path}")
        sys.exit(1)
