"""
Convert demo handout documents from Markdown to professionally formatted Word (.docx) files.
Run: python docs/demo/convert_to_word.py
Output: docs/demo/01_TO_GIOI_THIEU_SAN_PHAM.docx
        docs/demo/02_HUONG_DAN_SU_DUNG_NHANH.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
# STYLE / COLOR CONSTANTS
# ============================================================
BLUE_PRIMARY = RGBColor(0x1B, 0x4F, 0x72)    # Dark blue for headings
BLUE_ACCENT = RGBColor(0x21, 0x6F, 0xDB)     # Accent blue
RED_BETA = RGBColor(0xC0, 0x39, 0x2B)        # Red for beta warnings
GRAY_LIGHT = RGBColor(0x7F, 0x8C, 0x8D)      # Gray for footnotes
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEADER_BG = "1B4F72"                     # hex for table header shading
TABLE_ALT_BG = "EBF5FB"                        # hex for alternate row
FONT_NAME = "Times New Roman"
FONT_NAME_HEADING = "Arial"


def set_cell_shading(cell, hex_color):
    """Set background shading for a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Add thin borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_run(paragraph, text, bold=False, italic=False, size=12, color=BLACK, font=FONT_NAME):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = font
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font)
    return run


def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        p.space_before = Pt(18)
        p.space_after = Pt(6)
        make_run(p, text, bold=True, size=16, color=BLUE_PRIMARY, font=FONT_NAME_HEADING)
        # Add underline bar
        bar = doc.add_paragraph()
        bar_run = bar.add_run("─" * 80)
        bar_run.font.size = Pt(6)
        bar_run.font.color.rgb = RGBColor(0x21, 0x6F, 0xDB)
    elif level == 2:
        p.space_before = Pt(12)
        p.space_after = Pt(4)
        make_run(p, text, bold=True, size=13, color=BLUE_PRIMARY, font=FONT_NAME_HEADING)
    elif level == 3:
        p.space_before = Pt(8)
        p.space_after = Pt(2)
        make_run(p, text, bold=True, size=12, color=BLUE_ACCENT, font=FONT_NAME_HEADING)
    return p


def add_para(doc, text="", bold=False, italic=False, size=12, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.space_after = Pt(space_after)
    if text:
        make_run(p, text, bold=bold, italic=italic, size=size, color=color)
    return p


def add_bullet(doc, text, bold_prefix="", size=12):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        make_run(p, bold_prefix, bold=True, size=size)
        make_run(p, text, size=size)
    else:
        make_run(p, text, size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        make_run(p, h, bold=True, size=11, color=WHITE, font=FONT_NAME)
        set_cell_shading(cell, TABLE_HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            # Bold first column
            is_bold = (c_idx == 0)
            make_run(p, val, bold=is_bold, size=11)
            if r_idx % 2 == 1:
                set_cell_shading(cell, TABLE_ALT_BG)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_beta_watermark(doc):
    """Add a prominent beta/confidentiality notice at top."""
    # Warning box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(0)
    p.space_after = Pt(2)
    make_run(p, "⚠ TÀI LIỆU NỘI BỘ — KHÔNG LƯU HÀNH CÔNG KHAI ⚠", bold=True, size=12, color=RED_BETA, font=FONT_NAME_HEADING)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(12)
    make_run(p2, "Phần mềm đang trong giai đoạn thử nghiệm (Beta). Nội dung có thể thay đổi khi phát hành chính thức.", italic=True, size=10, color=RED_BETA)


def add_footer_beta(doc):
    """Add beta confidentiality reminder at the bottom."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, "─" * 60, size=8, color=GRAY_LIGHT)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(2)
    make_run(p2, "⚠ Tài liệu nội bộ — Phiên bản Beta, chưa phát hành chính thức.", italic=True, size=9, color=RED_BETA)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p3, "Không sao chép, phân phối khi chưa được sự đồng ý của đơn vị phát triển.", italic=True, size=9, color=GRAY_LIGHT)


def setup_page(doc, orientation="portrait"):
    """Configure page margins and orientation."""
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Set font for East Asian
    rPr = style.element.rPr
    if rPr is None:
        style.element.get_or_add_rPr()
        rPr = style.element.rPr
    rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


# ============================================================
# DOCUMENT 1: TỜ GIỚI THIỆU SẢN PHẨM
# ============================================================
def create_doc1():
    doc = Document()
    setup_page(doc)
    add_beta_watermark(doc)

    # ── TITLE ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    make_run(p, "PHẦN MỀM QUẢN LÝ VĂN BẢN THÔNG MINH", bold=True, size=14, color=BLUE_PRIMARY, font=FONT_NAME_HEADING)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.space_after = Pt(4)
    make_run(p2, "VANBANPLUS", bold=True, size=22, color=BLUE_ACCENT, font=FONT_NAME_HEADING)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(2)
    make_run(p3, "Phiên bản Beta — Tháng 02/2026", bold=True, size=11, color=GRAY_LIGHT)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.space_after = Pt(2)
    make_run(p4, "Đơn vị phát triển: Công ty TNHH Gia Kiệm Số", italic=True, size=11)
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.space_after = Pt(2)
    make_run(p5, "Website công ty: https://giakiemso.com", italic=True, size=11, color=BLUE_ACCENT)
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.space_after = Pt(12)
    make_run(p6, "Website phần mềm: https://vanbanplus.giakiemso.com", italic=True, size=11, color=BLUE_ACCENT)

    # ── SECTION: Giới thiệu ──
    add_heading_styled(doc, "Phần mềm VanBanPlus là gì?", level=1)
    add_para(doc, "VanBanPlus là công cụ hỗ trợ cá nhân dành cho cán bộ, công chức, viên chức — giúp quản lý, theo dõi văn bản và soạn thảo văn bản nhanh hơn nhờ ứng dụng trí tuệ nhân tạo (AI).")
    add_para(doc, "Phần mềm được cài đặt trên máy tính cá nhân. Dữ liệu lưu tại máy người dùng, đảm bảo an toàn thông tin.")
    add_para(doc, "⚠ Lưu ý: VanBanPlus là sổ tay công việc cá nhân — không thay thế hệ thống quản lý văn bản tập trung của cơ quan (eGov, iDesk, VNPT-iOffice...).", italic=True, size=11, color=GRAY_LIGHT)

    # ── Đối tượng ──
    add_heading_styled(doc, "Đối tượng sử dụng", level=1)
    add_table(doc,
        headers=["Đối tượng", "Dùng VanBanPlus để"],
        rows=[
            ["Cán bộ Văn phòng – Thống kê", "Nhập văn bản đến (scan PDF), AI soạn văn bản trả lời, theo dõi deadline, xuất báo cáo"],
            ["Lãnh đạo (Chủ tịch, Phó CT)", "Xem văn bản đến, ghi chú phân công, theo dõi tiến độ xử lý, duyệt nội dung"],
            ["Văn thư", "AI scan hàng loạt file PDF/ảnh, lưu sổ theo dõi, xuất danh sách văn bản"],
            ["Cán bộ chuyên môn", "Quản lý văn bản theo lĩnh vực, AI soạn chuyên ngành, lưu ảnh hiện trường"],
        ],
        col_widths=[5, 12]
    )
    add_para(doc, "Hỗ trợ hơn 70 loại hình cơ quan: UBND xã/huyện/tỉnh, HĐND, Đảng ủy, MTTQ, Hội Nông dân, Hội Phụ nữ, Đoàn Thanh niên, Hội Cựu chiến binh, các Sở ban ngành, Trường học, Trạm Y tế, Bệnh viện...", italic=True, size=11)

    # ── NHÓM 1 ──
    add_heading_styled(doc, "Nhóm 1 — Quản lý văn bản và cuộc họp", level=1)
    add_para(doc, "Không cần Internet. Hoạt động hoàn toàn trên máy tính.", italic=True, size=11, color=GRAY_LIGHT)

    add_table(doc,
        headers=["Chức năng", "Mô tả"],
        rows=[
            ["Quản lý VB đi/đến/nội bộ", "Tiếp nhận, phân loại, lưu trữ, tra cứu. Hỗ trợ đầy đủ 29 loại văn bản theo Điều 7, NĐ 30/2020/NĐ-CP"],
            ["Cây thư mục phân cấp", "Văn bản đến / Đi / Nội bộ / Lưu trữ. Tạo thư mục con không giới hạn"],
            ["Tìm kiếm và lọc", "Tìm theo từ khóa (hỗ trợ không dấu), lọc theo loại, ngày, người ký, hướng VB"],
            ["File đính kèm", "Đính kèm Word, PDF, Excel, ảnh. Đánh dấu file chính"],
            ["Sao văn bản", "Sao y, Sao lục, Trích sao — theo Điều 25–27, NĐ 30/2020"],
            ["Xuất file Word", "Đúng thể thức NĐ 30/2020 và TT 01/2011/TT-BNV. In trình ký được ngay"],
            ["Quản lý cuộc họp", "22 loại cuộc họp. Theo dõi người tham dự, nhiệm vụ, tiến độ. Xuất biên bản, kết luận"],
            ["Lịch tổng hợp", "Gom deadline văn bản, cuộc họp, nhiệm vụ vào 1 lịch. Phân biệt theo màu"],
            ["Thống kê", "So sánh theo tháng/quý/năm. Xuất Excel"],
            ["Album ảnh", "Lưu ảnh theo danh mục cơ quan. Upload/tải từ website"],
            ["Tra cứu pháp quy", "NĐ 30/2020 toàn văn, 38 Điều, 7 Chương, 6 Phụ lục. Tìm kiếm nhanh"],
            ["Sao lưu tự động", "Tự động sao lưu mỗi lần mở phần mềm. Khôi phục 1 bước"],
        ],
        col_widths=[5, 12]
    )

    # ── NHÓM 2 ──
    add_heading_styled(doc, "Nhóm 2 — Trí tuệ nhân tạo (AI)", level=1)
    add_para(doc, "Cần kết nối Internet khi sử dụng.", italic=True, size=11, color=GRAY_LIGHT)

    add_table(doc,
        headers=["Chức năng AI", "Mô tả", "Thời gian"],
        rows=[
            ["AI Scan trích xuất", "Chọn file PDF hoặc ảnh scan → AI tự đọc và điền 14 trường thông tin. Hỗ trợ xử lý nhiều file cùng lúc", "~15 giây/file"],
            ["AI Soạn thảo", "Chọn mẫu (41 mẫu có sẵn) → nhập thông tin chính → AI soạn bản nháp hoàn chỉnh, đúng thể thức NĐ 30/2020", "~20 giây"],
            ["AI Kiểm tra", "Kiểm tra 8 nhóm lỗi: chính tả, ngữ pháp, thể thức, căn cứ pháp lý, văn phong, nơi nhận, xung đột, gợi ý", "~15 giây"],
            ["AI Tham mưu", "Phân tích VB đến → đề xuất: ai xử lý, deadline, cần phúc đáp bằng gì, rủi ro nếu chậm trễ", "~15 giây"],
            ["AI Tóm tắt", "Tóm tắt VB dài (Nghị định, Thông tư) thành 10 mục: nội dung chính, đối tượng, thời hạn, tác động...", "~20 giây"],
            ["AI Báo cáo định kỳ", "Nhập số liệu thô → AI viết báo cáo hoàn chỉnh: kết quả, đánh giá, phương hướng. Tự tính tỷ lệ", "~30 giây"],
        ],
        col_widths=[4, 10, 3]
    )

    # ── Hiệu quả ──
    add_heading_styled(doc, "Hiệu quả ước tính", level=1)

    add_table(doc,
        headers=["Công việc", "Cách làm hiện tại", "Với VanBanPlus"],
        rows=[
            ["Nhập 1 VB từ file PDF", "10–15 phút", "2 phút"],
            ["Soạn 1 VB hành chính", "45–90 phút", "5–10 phút"],
            ["Kiểm tra VB trước trình ký", "Bị trả sửa 2–3 lần", "Giảm 70% lỗi"],
            ["Phân tích 1 VB đến", "30–60 phút", "3 phút"],
            ["Tóm tắt 1 VB dài", "2–3 giờ", "1 phút"],
            ["Viết 1 báo cáo định kỳ", "3–4 giờ", "15 phút"],
            ["Tra cứu VB cũ", "10–15 phút", "Vài giây"],
        ],
        col_widths=[5, 5, 5]
    )
    add_para(doc, "Ước tính cho 1 cán bộ Văn phòng: Tiết kiệm khoảng 3–5 giờ mỗi ngày.", bold=True, size=12, color=BLUE_PRIMARY)

    # ── Yêu cầu hệ thống ──
    add_heading_styled(doc, "Yêu cầu hệ thống", level=1)
    add_table(doc,
        headers=["Yêu cầu", "Chi tiết"],
        rows=[
            ["Hệ điều hành", "Windows 10 trở lên (64-bit)"],
            ["RAM", "Tối thiểu 4 GB"],
            ["Ổ cứng", "Khoảng 200 MB cho phần mềm"],
            ["Internet", "Không bắt buộc cho quản lý. Cần có cho chức năng AI"],
            ["Cài đặt", "1 lần, khoảng 5 phút"],
        ],
        col_widths=[5, 12]
    )

    # ── FAQ ──
    add_heading_styled(doc, "Câu hỏi thường gặp", level=1)

    faqs = [
        ("Dữ liệu có được bảo mật không?", "Dữ liệu lưu hoàn toàn trên máy tính của anh/chị, không lưu trên máy chủ bên ngoài."),
        ("Phần mềm có thay thế hệ thống eGov không?", "Không. VanBanPlus là công cụ hỗ trợ cá nhân, bổ sung cho hệ thống tập trung. Số văn bản chính thức vẫn do Văn thư cấp theo quy định."),
        ("AI có chính xác hoàn toàn không?", "AI đạt độ chính xác khoảng 95%. Cán bộ luôn cần kiểm tra lại kết quả trước khi sử dụng chính thức."),
        ("Máy tính bị hỏng thì dữ liệu có mất không?", 'Phần mềm tự động sao lưu mỗi lần khởi động. Có thể copy file sao lưu ra USB. Khi cài máy mới, chỉ cần nhấn "Khôi phục" là toàn bộ dữ liệu được phục hồi.'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        make_run(p, f"❓ {q}", bold=True, size=12, color=BLUE_PRIMARY)
        p2 = doc.add_paragraph()
        make_run(p2, f"→ {a}", size=11)
        p2.space_after = Pt(8)

    # ── FOOTER ──
    add_para(doc, "")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, "Công ty TNHH Gia Kiệm Số", bold=True, size=12, color=BLUE_PRIMARY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p2, "Website công ty: https://giakiemso.com", size=11, color=BLUE_ACCENT)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p3, "Website phần mềm: https://vanbanplus.giakiemso.com", size=11, color=BLUE_ACCENT)

    add_footer_beta(doc)

    out_path = os.path.join(SCRIPT_DIR, "01_TO_GIOI_THIEU_SAN_PHAM.docx")
    doc.save(out_path)
    print(f"✅ Đã tạo: {out_path}")
    return out_path


# ============================================================
# DOCUMENT 2: HƯỚNG DẪN SỬ DỤNG NHANH
# ============================================================
def create_doc2():
    doc = Document()
    setup_page(doc)
    add_beta_watermark(doc)

    # ── TITLE ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_before = Pt(12)
    make_run(p, "HƯỚNG DẪN SỬ DỤNG NHANH", bold=True, size=18, color=BLUE_PRIMARY, font=FONT_NAME_HEADING)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p2, "VANBANPLUS", bold=True, size=22, color=BLUE_ACCENT, font=FONT_NAME_HEADING)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.space_after = Pt(4)
    make_run(p3, "Tài liệu dành cho cán bộ, công chức sử dụng lần đầu", italic=True, size=11, color=GRAY_LIGHT)

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.space_after = Pt(16)
    make_run(p4, "Phiên bản Beta — Tháng 02/2026", size=11, color=GRAY_LIGHT)

    # ════════════════════════════════════════════════════════
    # 1. KHỞI ĐỘNG LẦN ĐẦU
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "1. Khởi động lần đầu", level=1)

    add_heading_styled(doc, "Bước 1: Mở phần mềm", level=3)
    add_bullet(doc, "Nhấp đúp biểu tượng VanBanPlus trên màn hình Desktop.")
    add_bullet(doc, "Lần đầu tiên, phần mềm hiện màn hình Thiết lập cơ quan.")

    add_heading_styled(doc, "Bước 2: Chọn loại cơ quan", level=3)
    add_bullet(doc, "Chọn đúng loại cơ quan nơi anh/chị công tác (UBND xã, Sở, Trường học, Đoàn thể...).")
    add_bullet(doc, "Phần mềm tự tạo cấu trúc thư mục phù hợp.")
    add_bullet(doc, 'Nhấn "Hoàn thành thiết lập".')

    add_heading_styled(doc, "Bước 3: Đăng nhập tài khoản", level=3)
    add_bullet(doc, "Nhấn vào nút Đăng nhập trên thanh bên trái.")
    add_bullet(doc, "Nhập email và mật khẩu tài khoản VanBanPlus.")
    add_bullet(doc, 'Nếu chưa có tài khoản → Nhấn "Đăng ký" → Nhập thông tin → Xác nhận.')

    add_para(doc, "⚠ Lưu ý: Cần đăng nhập để sử dụng các chức năng AI. Các chức năng quản lý văn bản cơ bản không cần đăng nhập.", italic=True, size=11, color=GRAY_LIGHT)

    # ════════════════════════════════════════════════════════
    # 2. AI SCAN
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "2. Nhập văn bản mới bằng AI Scan (cách nhanh nhất)", level=1)
    add_para(doc, "Đây là cách nhanh nhất để đưa văn bản vào hệ thống — đặc biệt phù hợp khi nhận file PDF, ảnh scan qua email hoặc Zalo.")

    steps = [
        'Tại trang Quản lý tài liệu → Nhấn nút "AI Scan" trên thanh công cụ.',
        'Nhấn "Chọn file" → Chọn file PDF hoặc ảnh scan cần nhập.',
        'Nhấn "Phân tích" → Chờ khoảng 10–15 giây.',
        "AI tự động điền đầy đủ các trường: Số ký hiệu, Ngày ban hành, Loại VB, Cơ quan, Trích yếu, Người ký, Nơi nhận, Căn cứ pháp lý...",
        "Kiểm tra lại thông tin AI đã điền → Chỉnh sửa nếu cần.",
        'Nhấn "Lưu".',
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph()
        make_run(p, f"Bước {i}: ", bold=True, size=12, color=BLUE_ACCENT)
        make_run(p, s, size=12)

    add_heading_styled(doc, "Nhập nhiều file cùng lúc:", level=3)
    add_bullet(doc, 'Nhấn "AI Scan" → Chọn nhiều file cùng lúc.')
    add_bullet(doc, "", bold_prefix="Tách riêng: ")
    # fix: add explanation after
    p = doc.add_paragraph(style='List Bullet')
    make_run(p, "Tách riêng: ", bold=True, size=12)
    make_run(p, "Mỗi file là 1 văn bản riêng biệt (dùng khi có nhiều VB khác nhau).", size=12)

    p = doc.add_paragraph(style='List Bullet')
    make_run(p, "Ghép trang: ", bold=True, size=12)
    make_run(p, "Nhiều file/ảnh ghép thành 1 VB (dùng khi 1 VB bị scan thành nhiều trang rời).", size=12)

    add_para(doc, "Định dạng file hỗ trợ: JPG, PNG, BMP, GIF, TIFF, WebP, PDF.", bold=True, size=11, color=GRAY_LIGHT)

    # ════════════════════════════════════════════════════════
    # 3. NHẬP THỦ CÔNG
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "3. Nhập văn bản thủ công", level=1)
    add_para(doc, 'Tại trang Quản lý tài liệu → Nhấn nút "Thêm". Điền các trường thông tin:')

    add_table(doc,
        headers=["Trường", "Ghi chú"],
        rows=[
            ["Hướng văn bản", "Chọn: Đến / Đi / Nội bộ"],
            ["Số ký hiệu", "Ví dụ: 123/QĐ-UBND"],
            ["Ngày ban hành", "Chọn ngày trên lịch"],
            ["Loại văn bản", "Chọn trong danh sách 29 loại"],
            ["Cơ quan ban hành", "Nhập tên cơ quan"],
            ["Trích yếu", "Nội dung tóm tắt"],
            ["Người ký", "Họ tên người ký"],
            ["Nơi nhận", "Các cơ quan/bộ phận nhận"],
            ["Mức độ khẩn", "Thường / Khẩn / Thượng khẩn / Hỏa tốc"],
            ["Căn cứ pháp lý", "Các văn bản được viện dẫn"],
            ["File đính kèm", "Đính kèm Word, PDF, Excel, ảnh"],
        ],
        col_widths=[5, 12]
    )
    add_para(doc, '→ Nhấn "Lưu". Phần mềm tự tạo Số đến tăng dần theo năm và Ký hiệu văn bản đúng format.', italic=True, size=11, color=GRAY_LIGHT)

    # ════════════════════════════════════════════════════════
    # 4. TÌM KIẾM
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "4. Tìm kiếm văn bản", level=1)

    add_heading_styled(doc, "Tìm nhanh:", level=3)
    add_bullet(doc, "Gõ từ khóa vào ô tìm kiếm phía trên danh sách (hỗ trợ tìm không dấu tiếng Việt).")
    add_bullet(doc, 'Ví dụ: Gõ "phong chong thien tai" → Tìm được "phòng chống thiên tai".')

    add_heading_styled(doc, "Lọc nhanh:", level=3)
    add_bullet(doc, "Nhấn các nút: Hôm nay / Tuần này / Tháng này để lọc theo thời gian.")

    add_heading_styled(doc, "Lọc nâng cao:", level=3)
    add_bullet(doc, "Lọc theo: Loại văn bản, Hướng (đi/đến), Khoảng ngày, Người ký, Số ký hiệu.")

    # ════════════════════════════════════════════════════════
    # 5. AI SOẠN THẢO
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "5. Soạn văn bản bằng AI", level=1)

    steps = [
        'Nhấn nút "Tạo văn bản AI" (trên Trang chủ hoặc thanh công cụ).',
        "Chọn mẫu văn bản — có sẵn 41 mẫu: Công văn, Quyết định, Báo cáo, Tờ trình, Kế hoạch, Thông báo...",
        'Nhập thông tin chính: nội dung yêu cầu, thành phần tham dự, người ký, chức danh.',
        'Nhấn "Tạo văn bản" → Chờ khoảng 15–20 giây.',
        "AI soạn bản nháp hoàn chỉnh, đúng thể thức theo NĐ 30/2020.",
        "Đọc lại, chỉnh sửa nếu cần.",
        'Nhấn "Lưu" → Nhấn "Xuất Word" để tạo file Word in trình ký.',
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph()
        make_run(p, f"Bước {i}: ", bold=True, size=12, color=BLUE_ACCENT)
        make_run(p, s, size=12)

    # ════════════════════════════════════════════════════════
    # 6. AI KIỂM TRA
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "6. Kiểm tra văn bản bằng AI", level=1)
    add_bullet(doc, "Mở văn bản cần kiểm tra (từ danh sách văn bản).")
    add_bullet(doc, 'Nhấn nút "AI Kiểm tra".')
    add_para(doc, "AI phân tích và hiển thị kết quả theo 3 mức độ:")

    add_table(doc,
        headers=["Mức độ", "Ý nghĩa", "Ví dụ"],
        rows=[
            ["🔴 Lỗi nghiêm trọng", "Phải sửa trước khi trình ký", "Thiếu căn cứ pháp lý, thiếu nơi nhận"],
            ["🟡 Cảnh báo", "Nên xem xét sửa", "Sai chính tả, câu quá dài, văn phong chưa phù hợp"],
            ["🟢 Gợi ý", "Có thể cải thiện", "Đề xuất bổ sung nội dung, cách diễn đạt tốt hơn"],
        ],
        col_widths=[4, 5, 8]
    )
    add_para(doc, "→ Nên kiểm tra AI trước mỗi lần trình ký để giảm tình trạng văn bản bị trả lại.", bold=True, italic=True, size=11, color=BLUE_PRIMARY)

    # ════════════════════════════════════════════════════════
    # 7. AI THAM MƯU
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "7. AI Tham mưu xử lý văn bản đến", level=1)
    add_para(doc, "Khi nhận được văn bản từ cấp trên và cần xác định phương án xử lý:")
    add_bullet(doc, "Mở văn bản đến cần phân tích.")
    add_bullet(doc, 'Nhấn nút "AI Tham mưu".')
    add_para(doc, "AI trả về phân tích gồm: tóm tắt nội dung, thời hạn xử lý, đề xuất người phụ trách, hình thức phúc đáp, căn cứ liên quan, cảnh báo rủi ro nếu chậm hạn, dàn ý dự thảo.")

    # ════════════════════════════════════════════════════════
    # 8. AI BÁO CÁO ĐỊNH KỲ
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "8. AI Báo cáo định kỳ", level=1)

    steps = [
        'Vào mục "AI Báo cáo định kỳ" trên thanh menu bên trái.',
        "Chọn kỳ báo cáo (Tuần / Tháng / Quý / Năm) và lĩnh vực.",
        'Nhập số liệu thô (ví dụ: "Thu ngân sách: 2,5 tỷ; Hộ nghèo giảm: 3 hộ; GPMB: đạt 85%").',
        "(Tùy chọn) Dán báo cáo kỳ trước để AI so sánh tỷ lệ tăng/giảm.",
        'Nhấn "Tạo báo cáo" → AI viết báo cáo hoàn chỉnh (kết quả, đánh giá, phương hướng).',
        "Đọc lại, chỉnh sửa → Xuất Word.",
    ]
    for i, s in enumerate(steps, 1):
        p = doc.add_paragraph()
        make_run(p, f"Bước {i}: ", bold=True, size=12, color=BLUE_ACCENT)
        make_run(p, s, size=12)

    # ════════════════════════════════════════════════════════
    # 9. CUỘC HỌP
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "9. Quản lý cuộc họp", level=1)
    add_bullet(doc, 'Vào mục Cuộc họp trên thanh menu.')
    add_bullet(doc, 'Nhấn "Thêm cuộc họp" → Điền: loại cuộc họp, thời gian, địa điểm, người tham dự, nội dung.')
    add_bullet(doc, 'Sau cuộc họp → Mở lại → Thêm biên bản, nhiệm vụ, tài liệu.')
    add_bullet(doc, 'Nhấn "Xuất Word" → Chọn: Biên bản / Kết luận / Báo cáo.')
    add_para(doc, "→ Phần mềm cảnh báo nếu tạo cuộc họp trùng lịch và theo dõi tiến độ nhiệm vụ sau họp.", italic=True, size=11, color=GRAY_LIGHT)

    # ════════════════════════════════════════════════════════
    # 10. SAO LƯU
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "10. Sao lưu và Khôi phục", level=1)

    add_heading_styled(doc, "Sao lưu:", level=3)
    add_bullet(doc, "Phần mềm tự động sao lưu mỗi lần khởi động.")
    add_bullet(doc, 'Sao lưu thủ công: Vào Sao lưu & Khôi phục → Nhấn "Sao lưu ngay".')
    add_bullet(doc, "Nên copy file sao lưu ra USB hoặc ổ cứng ngoài để phòng trường hợp máy hỏng.")

    add_heading_styled(doc, "Khôi phục:", level=3)
    add_bullet(doc, 'Vào Sao lưu & Khôi phục → Nhấn "Khôi phục" → Chọn file sao lưu → Dữ liệu phục hồi hoàn toàn.')

    # ════════════════════════════════════════════════════════
    # 11. PHÍM TẮT
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "11. Phím tắt thông dụng", level=1)
    add_table(doc,
        headers=["Phím tắt", "Chức năng"],
        rows=[
            ["Ctrl + N", "Thêm văn bản mới"],
            ["Ctrl + F", "Tìm kiếm"],
            ["Delete", "Xóa văn bản đang chọn"],
            ["F5", "Làm mới danh sách"],
            ["F1", "Mở trang Trợ giúp"],
        ],
        col_widths=[5, 12]
    )

    # ════════════════════════════════════════════════════════
    # 12. HỖ TRỢ
    # ════════════════════════════════════════════════════════
    add_heading_styled(doc, "12. Cần hỗ trợ?", level=1)
    add_bullet(doc, "Nhấn phím F1 trong phần mềm để mở trang Trợ giúp chi tiết.")
    add_para(doc, "Liên hệ đội ngũ kỹ thuật:")
    add_bullet(doc, "Công ty TNHH Gia Kiệm Số")
    add_bullet(doc, "Website công ty: https://giakiemso.com")
    add_bullet(doc, "Website phần mềm: https://vanbanplus.giakiemso.com")

    add_footer_beta(doc)

    out_path = os.path.join(SCRIPT_DIR, "02_HUONG_DAN_SU_DUNG_NHANH.docx")
    doc.save(out_path)
    print(f"✅ Đã tạo: {out_path}")
    return out_path


# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("=" * 50)
    print("Đang tạo tài liệu Word cho buổi demo...")
    print("=" * 50)
    create_doc1()
    create_doc2()
    print("\n✅ Hoàn thành! 2 file .docx đã sẵn sàng trong thư mục docs/demo/")
